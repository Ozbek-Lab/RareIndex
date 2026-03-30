from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile

from django.conf import settings


@dataclass(frozen=True)
class DocxTemplateChoice:
    path: Path | None
    reason: str


def _reports_dir() -> Path:
    return Path(settings.BASE_DIR) / "reports"


def resolve_docx_template_path(template_location: str | None) -> Path | None:
    if not template_location:
        return None

    raw = Path(template_location)
    if raw.is_absolute():
        return raw if raw.exists() else None

    candidate = Path(settings.BASE_DIR) / template_location
    if candidate.exists():
        return candidate

    reports_candidate = _reports_dir() / template_location
    if reports_candidate.exists():
        return reports_candidate

    return None


def choose_docx_template_for_test_type(test_type_name: str | None, *, negative: bool) -> DocxTemplateChoice:
    """
    Pick a DOCX template from the repo's `reports/` directory.

    Heuristic (simple + robust):
    - Prefer templates that match the test type token (wes/wgs/sanger/cma/rna/...) in filename.
    - Prefer negative/negatif templates when negative=True.
    - Fall back to the first .docx found; if none exist, return None (caller will create a blank doc).
    """
    reports_dir = _reports_dir()
    if not reports_dir.exists():
        return DocxTemplateChoice(path=None, reason=f"{reports_dir} missing")

    candidates = sorted([p for p in reports_dir.glob("*.docx") if p.is_file()])
    if not candidates:
        return DocxTemplateChoice(path=None, reason=f"no .docx in {reports_dir}")

    name_l = (test_type_name or "").strip().lower()
    tokens: list[str] = []
    if "wes" in name_l:
        tokens.append("wes")
    if "wgs" in name_l:
        tokens.append("wgs")
    if "sanger" in name_l:
        tokens.append("sanger")
    if "cma" in name_l or "array" in name_l:
        tokens.append("cma")
    if "rna" in name_l:
        tokens.append("rna")
    if not tokens and name_l:
        tokens.append(name_l.split()[0])

    def score(p: Path) -> tuple[int, int, int]:
        fn = p.name.lower()
        neg_score = 0
        if negative:
            neg_score = 2 if any(t in fn for t in ("negatif", "negative", "neg")) else 0
        else:
            neg_score = 1 if not any(t in fn for t in ("negatif", "negative", "neg")) else 0

        token_score = 0
        for t in tokens:
            if t and t in fn:
                token_score = 3
                break

        length_score = -len(fn)
        return (token_score, neg_score, length_score)

    best = max(candidates, key=score)
    return DocxTemplateChoice(path=best, reason=f"best match among {len(candidates)} candidates")


def build_docx_report_bytes(
    *,
    template_path: Path | None,
    title: str,
    variants_rows: Iterable[dict],
    negative: bool,
    placeholders: dict[str, str] | None = None,
    rich_text_blocks: dict[str, list[list[str]]] | None = None,
) -> bytes:
    from docx import Document

    if template_path and template_path.exists():
        if placeholders:
            rendered = _render_docx_template(
                template_path,
                placeholders,
                skip_keys=set((rich_text_blocks or {}).keys()),
            )
            if rich_text_blocks:
                return _apply_rich_text_blocks(rendered, rich_text_blocks)
            return rendered
        doc = Document(str(template_path))
    else:
        doc = Document()

    # Spacer
    doc.add_paragraph("")
    # Some legacy templates may not define "Heading 2"; fall back to a bold paragraph.
    try:
        doc.add_heading(title, level=2)
    except KeyError:
        p = doc.add_paragraph(title)
        run = p.runs[0] if p.runs else p.add_run(title)
        run.bold = True

    if negative:
        doc.add_paragraph("No clinically relevant variants were reported.")
    else:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Location"
        hdr[1].text = "Zygosity"
        hdr[2].text = "Type"
        hdr[3].text = "Genes"

        for row in variants_rows:
            cells = table.add_row().cells
            cells[0].text = str(row.get("location") or "")
            cells[1].text = str(row.get("zygosity") or "")
            cells[2].text = str(row.get("type") or "")
            cells[3].text = str(row.get("genes") or "")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_template(
    template_path: Path,
    placeholders: dict[str, str],
    *,
    skip_keys: set[str] | None = None,
) -> bytes:
    """Replace placeholder tokens inside DOCX XML parts."""
    xml_members = [
        "word/document.xml",
        "word/header1.xml",
        "word/header2.xml",
        "word/footer1.xml",
        "word/footer2.xml",
    ]

    def encode_value(value: str) -> str:
        escaped = escape(value or "")
        # DOCX placeholders in this project live inside single text nodes,
        # so we can safely expand newlines into explicit line breaks.
        return escaped.replace("\n", '</w:t><w:br/><w:t xml:space="preserve">')

    rendered = BytesIO()
    with ZipFile(template_path, "r") as zin, ZipFile(rendered, "w", compression=ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in xml_members:
                text = data.decode("utf-8")
                for key, value in placeholders.items():
                    if skip_keys and key in skip_keys:
                        continue
                    text = text.replace(f"{{{{{key}}}}}", encode_value(value))
                data = text.encode("utf-8")
            zout.writestr(item, data)

    return rendered.getvalue()


def _apply_rich_text_blocks(docx_bytes: bytes, rich_text_blocks: dict[str, list[list[str]]]) -> bytes:
    from docx import Document

    doc = Document(BytesIO(docx_bytes))
    token_map = {f"{{{{{key}}}}}": entries for key, entries in rich_text_blocks.items()}

    for paragraph in _iter_all_paragraphs(doc):
        for token, entries in token_map.items():
            if token in paragraph.text:
                _replace_paragraph_with_rich_text(paragraph, entries, font_size_pt=5.5)
                break

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _iter_all_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def _replace_paragraph_with_rich_text(paragraph, entries: list[list[str]], font_size_pt: float | None = None) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    normalized_entries = [
        [line.strip() for line in entry if line and line.strip()]
        for entry in entries
    ]
    normalized_entries = [entry for entry in normalized_entries if entry]

    for entry_index, lines in enumerate(normalized_entries):
        for line_index, line in enumerate(lines):
            run = paragraph.add_run(line)
            run.bold = line_index == 0
            if font_size_pt is not None:
                run.font.size = Pt(font_size_pt)
            if line_index < len(lines) - 1:
                run.add_break()
        if entry_index < len(normalized_entries) - 1:
            paragraph.add_run().add_break()
            paragraph.add_run().add_break()


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
